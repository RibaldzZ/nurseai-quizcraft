import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Style config
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ===== COVER PAGE =====
for _ in range(6):
    doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('NurseAI QuizCraft')
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0x1E, 0x40, 0x73)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('AI-Powered NGN Question Generator\nfor Nurse Educators')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x47, 0x6D, 0xA8)

doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Project Documentation')
run.font.size = Pt(14)
run.italic = True

doc.add_paragraph('')
doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Version 1.0\nJune 2026\n\nBuilt by NursingAI\nPowered by Google Gemini AI')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ===== TABLE OF CONTENTS =====
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Project Title',
    '2. Overview / System Features',
    '3. Architecture Design',
    '4. Objectives',
    '5. Scope',
    '6. Technical Requirements',
    '7. How to Use / System Workflow',
    '8. Troubleshooting Guide',
]
for item in toc_items:
    doc.add_paragraph(item)

doc.add_page_break()

# ===== SECTION 1 =====
doc.add_heading('1. Project Title', level=1)
doc.add_paragraph('NurseAI QuizCraft')

items = [
    ('Tagline: ', 'AI-Powered NGN Question Generator for Nurse Educators'),
    ('Platform: ', 'Web Application (Next.js 15.2.9)'),
    ('AI Engine: ', 'Google Gemini 2.5 Flash Lite'),
    ('Target Users: ', 'Nurse Faculty / Instructors, Nursing Education Programs'),
]
for label, value in items:
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.bold = True
    p.add_run(value)

doc.add_page_break()

# ===== SECTION 2 =====
doc.add_heading('2. Overview / System Features', level=1)
doc.add_paragraph(
    'NurseAI QuizCraft is a web-based AI application designed to help nurse educators '
    'create Next Generation NCLEX (NGN)-style exam questions quickly and efficiently. '
    'Using Google Gemini AI, the system generates clinically accurate questions with '
    'realistic patient scenarios, evidence-based rationales, and customizable difficulty levels.'
)

doc.add_heading('Core Features', level=2)
features = [
    ('Multi-Format NGN Question Generation',
     'Supports five question types: Multiple Choice, Case Study, Bow-Tie (Trend), Cloze, and Matrix.'),
    ('AI-Powered Content Creation',
     'Google Gemini 2.5 Flash Lite generates questions with realistic patient scenarios and rationales.'),
    ('Difficulty Customization',
     'Three levels (Easy, Moderate, Difficult) to match learner progression.'),
    ('Adjustable Quantity',
     'Generate 1-5 questions per request.'),
    ('Integrated Answer Key',
     'Toggle answer key on/off with detailed rationales.'),
    ('Copy-to-Clipboard Export',
     'Export all questions in clean text format.'),
    ('Responsive Web Design',
     'Works on desktop, tablet, and mobile.'),
    ('Zero Database Requirement',
     'No user accounts or data storage needed.'),
]
for title, desc in features:
    p = doc.add_paragraph()
    run = p.add_run(title + ': ')
    run.bold = True
    p.add_run(desc)

doc.add_page_break()

# ===== SECTION 3 =====
doc.add_heading('3. Architecture Design', level=1)
doc.add_paragraph(
    'The application follows a modern, lightweight client-server architecture optimized for '
    'rapid development and one-click deployment on Vercel.'
)

doc.add_heading('Data Flow', level=2)
steps = [
    'Step 1: User fills out form (topic, question type, difficulty, count)',
    'Step 2: Frontend sends POST request to /api/generate endpoint',
    'Step 3: API route validates input and calls Google Gemini API with structured prompt',
    'Step 4: Gemini generates NGN-style questions and returns JSON response',
    'Step 5: Response normalizer transforms AI output to match expected TypeScript types',
    'Step 6: API route returns parsed questions to frontend',
    'Step 7: Frontend renders questions as interactive cards with toggle answer key',
    'Step 8: User can copy all questions to clipboard',
]
for step in steps:
    doc.add_paragraph(step, style='List Bullet')

doc.add_heading('Technology Stack', level=2)
stack = [
    ('Frontend', 'Next.js 15.2.9 (App Router), React 19, TypeScript 5, Tailwind CSS v4'),
    ('Backend', 'Next.js API Routes (serverless functions)'),
    ('AI Engine', 'Google Gemini 2.5 Flash Lite via @google/generative-ai SDK'),
    ('Deployment', 'Vercel (serverless, edge-ready)'),
    ('Version Control', 'Git, GitHub'),
]
for title, desc in stack:
    p = doc.add_paragraph()
    run = p.add_run(title + ': ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('Architecture Diagram', level=2)
diagram_lines = [
    '+-------------------+       +-------------------+       +------------------------+',
    '|                   |       |                   |       |                        |',
    '|  Next.js Frontend | ----> |  /api/generate    | ----> |  Google Gemini 2.5     |',
    '|  (React 19,       |       |  (Serverless fn)  |       |  Flash Lite API        |',
    '|   Tailwind v4)    | <---- |                   | <---- |                        |',
    '+-------------------+       +-------------------+       +------------------------+',
]
for line in diagram_lines:
    p = doc.add_paragraph()
    run = p.add_run(line)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)

doc.add_page_break()

# ===== SECTION 4 =====
doc.add_heading('4. Objectives', level=1)
doc.add_heading('Primary Objective', level=2)
doc.add_paragraph(
    'To provide nurse educators with an AI-powered tool that rapidly generates high-quality, '
    'clinically accurate NGN-style exam questions, reducing question creation time from 30+ minutes '
    'per item to under 30 seconds.'
)

doc.add_heading('Specific Objectives', level=2)
objectives = [
    'Reduce faculty workload in exam preparation by automating question generation.',
    'Improve alignment of nursing exam items with Next Generation NCLEX (NGN) formats.',
    'Ensure clinical accuracy through AI models trained on evidence-based nursing content.',
    'Provide detailed rationales to support both formative and summative assessment.',
    'Enable faculty to customize difficulty levels based on learner progression.',
    'Support multiple NGN question types to prepare students for the current NCLEX format.',
    'Eliminate the need for complex software installations.',
    'Facilitate rapid deployment on Vercel with zero infrastructure management.',
]
for obj in objectives:
    doc.add_paragraph(obj, style='List Bullet')

doc.add_page_break()

# ===== SECTION 5 =====
doc.add_heading('5. Scope', level=1)

doc.add_heading('In Scope (MVP v1.0)', level=2)
in_scope = [
    'Single-page web application accessible via modern web browsers.',
    'Generation of five NGN question types: Multiple Choice, Case Study, Bow-Tie, Cloze, Matrix.',
    'Topic-based question generation with customizable difficulty levels.',
    'Display of questions with toggleable answer key and rationale.',
    'Copy-to-clipboard export of all generated questions.',
    'Integration with Google Gemini 2.5 Flash Lite API.',
    'Clean, responsive UI optimized for desktop and mobile.',
    'One-click deployment on Vercel.',
    'Response normalizer to handle AI output formatting variations.',
]
for item in in_scope:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Out of Scope (Future Versions)', level=2)
out_scope = [
    'User authentication and role-based access control.',
    'Database storage for question banks and history.',
    'Bulk upload / import of topics.',
    'Export to LMS formats (QTI, CSV, IMS).',
    'Image/media embedding in questions.',
    'Student-facing test delivery and scoring.',
    'Analytics dashboard for faculty.',
    'Multi-language support.',
    'Offline mode.',
]
for item in out_scope:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ===== SECTION 6 =====
doc.add_heading('6. Technical Requirements', level=1)

doc.add_heading('Hardware Requirements', level=2)
doc.add_paragraph(
    'No special hardware required. Application runs entirely in the cloud. '
    'End users only need a device with internet access and a modern web browser.'
)

doc.add_heading('Software Requirements', level=2)
reqs = [
    ('Runtime', 'Node.js 18+ (local dev only)'),
    ('Package Manager', 'npm or yarn'),
    ('Framework', 'Next.js 15.2.9 (App Router)'),
    ('Language', 'TypeScript 5.x'),
    ('Styling', 'Tailwind CSS v4'),
    ('AI SDK', '@google/generative-ai v0.24+'),
    ('Browser Support', 'Chrome 90+, Firefox 90+, Safari 15+, Edge 90+'),
    ('Deployment Platform', 'Vercel (recommended) or any Node.js hosting'),
]
for title, desc in reqs:
    p = doc.add_paragraph()
    run = p.add_run(title + ': ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('API Requirements', level=2)
doc.add_paragraph('Google Gemini API key (free tier at aistudio.google.com)', style='List Bullet')
doc.add_paragraph('Gemini 2.5 Flash Lite model recommended (stable, cost-effective)', style='List Bullet')
doc.add_paragraph('Free tier provides sufficient quota for development and testing', style='List Bullet')

doc.add_heading('Environment Variables', level=2)
p = doc.add_paragraph()
run = p.add_run('GEMINI_API_KEY')
run.font.name = 'Courier New'
p.add_run(': Your Google Gemini API key (required)')

doc.add_heading('Deployment Requirements (Vercel)', level=2)
doc.add_paragraph('GitHub repository to host the source code.', style='List Bullet')
doc.add_paragraph('Vercel account (free tier sufficient).', style='List Bullet')
doc.add_paragraph('Environment variable GEMINI_API_KEY configured in Vercel dashboard.', style='List Bullet')
doc.add_paragraph('Node.js 18.x selected in Vercel project settings.', style='List Bullet')

doc.add_page_break()

# ===== SECTION 7 =====
doc.add_heading('7. How to Use / System Workflow', level=1)

doc.add_heading('Workflow Diagram', level=2)
p = doc.add_paragraph()
run = p.add_run('START -> Enter Topic -> Select Question Type -> Set Difficulty -> Pick Quantity -> Click Generate -> AI Creates Questions -> Review -> Toggle Answer Key -> Copy All -> END')
run.font.size = Pt(10)

doc.add_heading('Step-by-Step Guide', level=2)

guide = [
    ('Step 1: Access the Application',
     'Open the web app in any modern browser. No login or registration required.'),
    ('Step 2: Enter a Topic',
     'Type a nursing topic (e.g., Fluid and Electrolytes, Cardiac Nursing, Pharmacology).'),
    ('Step 3: Select Question Type',
     'Choose from five NGN formats: Multiple Choice, Case Study, Bow-Tie, Cloze, or Matrix.'),
    ('Step 4: Set Difficulty Level',
     'Select Easy, Moderate, or Difficult to match your students progression.'),
    ('Step 5: Choose Quantity',
     'Select how many questions to generate (1-5 per request).'),
    ('Step 6: Generate Questions',
     'Click the Generate Questions button. AI processing takes 5-15 seconds.'),
    ('Step 7: Review Questions',
     'Each question appears in an interactive card. Scroll through all generated items.'),
    ('Step 8: Toggle Answer Key',
     'Click Show Answer Key to reveal correct answers and detailed rationales.'),
    ('Step 9: Copy to Clipboard',
     'Click Copy All to export all questions in clean text format.'),
]
for title, desc in guide:
    doc.add_heading(title, level=3)
    doc.add_paragraph(desc)

doc.add_page_break()

# ===== SECTION 8 =====
doc.add_heading('8. Troubleshooting Guide', level=1)

troubleshoot = [
    ('Next.js Outdated Error',
     'Update Next.js with: npm install next@latest. Project uses 15.2.9.'),
    ('API Key Not Set',
     'Ensure GEMINI_API_KEY is in .env.local (local) or Vercel Environment Variables.'),
    ('429 Too Many Requests',
     'Free tier quota exhausted. Wait 1 minute or create a new Google AI project.'),
    ('503 Service Unavailable',
     'Model overloaded. App uses gemini-2.5-flash-lite for best reliability.'),
    ('AI Response Format Errors',
     'The built-in response normalizer handles most formatting issues automatically.'),
    ('Build Failures on Vercel',
     'Run npm run build locally first to catch TypeScript/ESLint errors.'),
    ('App Loads but Shows Blank Screen',
     'Check browser console (F12). Common cause: missing environment variable.'),
]
for title, desc in troubleshoot:
    p = doc.add_paragraph()
    run = p.add_run(title + ': ')
    run.bold = True
    p.add_run(desc)

# Save
output_path = r'C:\Users\ribaldz\Desktop\NursingAI_Project\NurseAI_QuizCraft_Documentation.docx'
doc.save(output_path)
print('Document saved successfully!')
